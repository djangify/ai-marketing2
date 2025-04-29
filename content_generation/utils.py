# content_generation/utils.py
import io
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

def generate_docx(title, content):
    """
    Generate a DOCX document from content with proper Unicode handling
    
    Args:
        title (str): Title for the document
        content (str): Content text
        
    Returns:
        BytesIO: Document as a byte stream
    """
    document = Document()
    
    # Add proper font settings
    font_name = 'Arial'
    styles = document.styles
    style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = font_name
    font.size = Pt(11)
    
    # Set style for the document
    for p in document.paragraphs:
        p.style = 'Normal'
    
    # Add title
    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = font_name
    
    # Add content - handle line breaks properly
    paragraphs = content.split('\n\n')
    for paragraph_text in paragraphs:
        if paragraph_text.strip():
            p = document.add_paragraph(style='Normal')
            p.add_run(paragraph_text.strip())
    
    # Save to memory stream
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
